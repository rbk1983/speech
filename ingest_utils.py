# ingest_utils.py — helpers for adding new content (URL, files, audio/video) and rebuilding index
import os, io, json, re, datetime as _dt
import pandas as pd
import numpy as np
import faiss
import requests
from bs4 import BeautifulSoup
from readability import Document
import pdfplumber
from docx import Document as DocxDocument
from openai import OpenAI

EMBED_MODEL = "text-embedding-3-large"
_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def _clean_text(t: str) -> str:
    t = re.sub(r'\s+', ' ', t or '').strip()
    # Remove boilerplate lines seen on IMF pages
    junk = [
        "The IMF Press Center is a password-protected site for working journalists",
        "Sign up to receive free e-mail notices"
    ]
    for j in junk:
        t = t.replace(j, ' ')
    return t.strip()

def extract_from_url(url: str) -> dict:
    try:
        r = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
        r.raise_for_status()
        doc = Document(r.text)
        title = doc.short_title() or ""
        html = doc.summary()
        soup = BeautifulSoup(html, "html.parser")
        text = _clean_text(soup.get_text(separator=" "))
        # Try meta date
        date = ""
        s2 = BeautifulSoup(r.text, "html.parser")
        for sel in [
            {"name":"meta","attrs":{"property":"article:published_time"}},
            {"name":"meta","attrs":{"name":"date"}},
            {"name":"meta","attrs":{"itemprop":"datePublished"}},
        ]:
            tag = s2.find(**sel)
            if tag and (tag.get("content") or tag.get("value")):
                date = (tag.get("content") or tag.get("value")).split("T")[0]
                break
        return {"title": title, "text": text, "date": date or ""}
    except Exception:
        return {"title":"", "text":"", "date":""}

def extract_from_pdf(file_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages = [p.extract_text() or "" for p in pdf.pages]
    return _clean_text(" ".join(pages))

def extract_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paras = [p.text for p in doc.paragraphs if p.text]
    return _clean_text(" ".join(paras))

def extract_from_txt(file_bytes: bytes) -> str:
    return _clean_text(file_bytes.decode("utf-8", errors="ignore"))

def transcribe_media(file_bytes: bytes, filename: str) -> str:
    # Send bytes directly to OpenAI Whisper; supports common audio/video containers.
    f = io.BytesIO(file_bytes)
    f.name = filename
    try:
        tr = _client.audio.transcriptions.create(model="whisper-1", file=f)
        return _clean_text(tr.text or "")
    except Exception as e:
        raise RuntimeError(f"Transcription failed: {e}")

def normalize_record(title, date, link, location, transcript, themes=""):
    # themes: comma-separated string -> list
    theme_list = [t.strip() for t in (themes or "").split(",") if t.strip()]
    return {
        "title": title or "(untitled)",
        "date": date,
        "link": link or "",
        "location": location or "",
        "transcript": transcript or "",
        "themes": theme_list
    }

def _chunk(text, max_tokens_chars=1800, overlap_chars=300):
    # Approx by characters to avoid tokenizer dependency
    text = text or ""
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        j = min(n, i + max_tokens_chars)
        chunk = text[i:j]
        chunks.append(chunk)
        i = j - overlap_chars
        if i <= 0: i = j
    return chunks

def _embed(texts):
    embs = _client.embeddings.create(model=EMBED_MODEL, input=texts).data
    return np.array([e.embedding for e in embs], dtype="float32")

def add_records_and_rebuild(new_records: list) -> tuple:
    """
    Append new records to speeches_data.pkl (or reclassified), dedupe, rebuild FAISS.
    Returns (ok: bool, message: str)
    """
    try:
        if os.path.exists("speeches_data_reclassified.pkl"):
            df = pd.read_pickle("speeches_data_reclassified.pkl")
        elif os.path.exists("speeches_data.pkl"):
            df = pd.read_pickle("speeches_data.pkl")
        else:
            df = pd.DataFrame(columns=["title","date","link","location","transcript","themes"])

        add_df = pd.DataFrame(new_records)
        add_df["date"] = pd.to_datetime(add_df["date"], errors="coerce").dt.date.astype(str)
        df["date"] = pd.to_datetime(df["date"], errors="coerce").dt.date.astype(str)

        combined = pd.concat([df, add_df], ignore_index=True)
        combined["key"] = combined["title"].fillna("") + "||" + combined["date"].fillna("") + "||" + combined["link"].fillna("")
        combined = combined.drop_duplicates(subset=["key"]).drop(columns=["key"])

        # Save back
        target = "speeches_data_reclassified.pkl" if os.path.exists("speeches_data_reclassified.pkl") else "speeches_data.pkl"
        combined.to_pickle(target)

        # Rebuild index
        metas, chunks = [], []
        for _, row in combined.iterrows():
            text = (row.get("transcript") or "").strip()
            if not text: continue
            parts = _chunk(text)
            for k, part in enumerate(parts):
                metas.append({
                    "title": row.get("title"),
                    "date": row.get("date"),
                    "link": row.get("link"),
                    "location": row.get("location",""),
                    "themes": row.get("themes") or [],
                    "chunk": k
                })
                chunks.append(part)

        if not chunks:
            return False, "No text to index."
        X = _embed(chunks)
        index = faiss.IndexFlatIP(X.shape[1])
        faiss.normalize_L2(X)
        index.add(X)
        faiss.write_index(index, "index.faiss")
        with open("meta.json","w",encoding="utf-8") as f:
            json.dump({"metas": metas}, f, ensure_ascii=False)
        with open("chunks.json","w",encoding="utf-8") as f:
            json.dump({"chunks": chunks}, f, ensure_ascii=False)
        return True, f"Added {len(add_df)} item(s). Corpus now has {len(combined)} records; index has {len(chunks)} chunks."
    except Exception as e:
        return False, f"Ingestion failed: {e}"
